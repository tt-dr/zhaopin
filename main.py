import re
import asyncio
import random
import os
import json
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm.asyncio import tqdm_asyncio

# ==================== 全局配置常量 ====================
START_URL = "https://www.zhaopin.com/sou/jl530/kwFT8NTN2RH58MG/p1" #起始的URL
MAX_PAGES = 1  #可修改，一次性爬去的最大页数
START_PAGE = 1  #可修改：从第几页开始爬取
MAX_CONCURRENT = 10 #如果觉得满可以加一点，这个是同时挂载的网页数，如果一次挂十个电脑变得特别卡，就调小一点
SLEEP_MIN = 0.4
SLEEP_MAX = 0.8

PROXY = None
WORD_FILENAME = "新北京岗位信息.docx" #写入的word文件名，这个默认写到同文件夹下面，这个会自动创建的
JSON_FILENAME = "新北京岗位信息.json" #写入的json文件名，这个后面进行数据清洗的时候会用


RESUME_FILE = "../resume_state.json"
JOB_TITLE_SELECTORS = ["h3.summary-plane__title", "h1.job-title", "h3.title", ".job-name"]
SALARY_SELECTORS = [".summary-plane__salary", ".job-salary", ".salary"]
COMPANY_NAME_SELECTORS = [".company__title a", ".company-name a", ".job-company__name"]
DETAIL_CONTENT_SELECTORS = [".description__detail-content", ".job-detail-content", ".job-description", ".describtion"]
ADDRESS_SELECTORS = [".job-address__content-text", ".work-add", ".job-location"]
BASIC_INFO_SELECTORS = [".summary-plane__info li", ".job-basic-info li"]
COMPANY_INFO_SELECTORS = [".company__info", ".job-company__info", ".company-details"]
JOB_LINK_SELECTORS = [
    ".joblist-box__item .jobinfo__top a",
    ".joblist-box__item .jobinfo_top a",
    ".job-item a",
    ".job-card a",
    ".job-list .job-title a"
]

LOCATION_PATTERN = r'([^，,]+?市[^，,]*?区?)'
EXPERIENCE_PATTERNS = [r'(\d-\d年|\d+年以上|经验不限|应届)', r'(\d+年经验)']
EDUCATION_PATTERN = r'(本科|大专|硕士|博士|高中|中专|不限)'
COMPANY_SIZE_PATTERN = r'(少于50人|50-100人|100-500人|500-1000人|1000人以上)'

JOB_DESC_PATTERNS = [
    r'([职责描述工作内容岗位职责主要职责]+)[:：]?(.*?)([任职要求岗位要求任职资格要求]+)[:：]?(.*)',
    r'([任职要求岗位要求任职资格要求]+)[:：]?(.*?)([职责描述工作内容岗位职责主要职责]+)[:：]?(.*)',
    r'([岗位职责工作职责职责描述]+)(.*?)([任职要求岗位要求任职资格]+)(.*)',
    r'([任职要求岗位要求任职资格]+)(.*?)([岗位职责工作职责职责描述]+)(.*)'
]


class ZhaopinScraper:
    def __init__(self, start_url, max_pages=1):
        self.start_url = start_url
        self.max_pages = max_pages
        self.session_data = []
        self.word_initialized = False
        self.resume_file = RESUME_FILE
        self.state = self.load_resume_state()
        self.crawled_pages = set(self.state.get("completed_pages", []))
        self.seen_links = set(self.state.get("seen_links", []))

        if os.path.exists(JSON_FILENAME):
            try:
                with open(JSON_FILENAME, 'r', encoding='utf-8') as f:
                    existing = json.load(f)
                    self.session_data.extend(existing)
                    for job in existing:
                        link = job.get("详情链接")
                        if link:
                            self.seen_links.add(link)
            except Exception as e:
                print(f"载入已有 JSON 失败: {e}")

    def load_resume_state(self):
        if os.path.exists(self.resume_file):
            try:
                with open(self.resume_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                pass
        return {"completed_pages": [], "seen_links": []}

    def save_resume_state(self):
        state = {
            "completed_pages": list(self.crawled_pages),
            "seen_links": list(self.seen_links)
        }
        with open(self.resume_file, 'w', encoding='utf-8') as f:
            json.dump(state, f, ensure_ascii=False, indent=2)

    def init_word_doc(self, filename):
        doc = Document()
        title = doc.add_heading("智联招聘岗位信息", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(filename)
        self.word_initialized = True
        print(f"已创建文档：{filename}")

    def append_page_to_word(self, jobs, page_num, filename=WORD_FILENAME):
        if not self.word_initialized and not os.path.exists(filename):
            self.init_word_doc(filename)

        doc = Document(filename)
        doc.add_paragraph(f"\n{'='*60}")
        doc.add_paragraph(f"第 {page_num} 页职位信息（共 {len(jobs)} 个）")
        doc.add_paragraph(f"{'='*60}\n")

        for i, job in enumerate(jobs, 1):
            doc.add_paragraph("-" * 50)
            job_title = doc.add_heading(f"岗位 {i}: {job.get('岗位名称', '未知')}", level=1)

            fields = [
                ("薪资", "薪资"),
                ("工作地点", "工作地点"),
                ("公司名称", "公司名称"),
                ("经验要求", "经验要求"),
                ("学历要求", "学历要求"),
                ("公司规模", "公司规模"),
                ("公司行业", "公司行业"),
                ("工作职责", "工作职责"),
                ("任职要求", "任职要求")
            ]

            for display_name, key in fields:
                value = job.get(key, "")
                if value and value not in {"未获取到", "获取失败"}:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{display_name}：")
                    run.bold = True
                    p.add_run(str(value))

        doc.save(filename)
        print(f"第 {page_num} 页已写入文档：{filename}")

    async def safe_text(self, page, selectors):
        for selector in selectors:
            try:
                element = page.locator(selector)
                if await element.count() > 0:
                    text = await element.inner_text()
                    if text and text.strip():
                        return text.strip()
            except:
                continue
        return ""

    async def get_basic_info(self, page):
        info_dict = {}
        info_items = await page.locator(", ".join(BASIC_INFO_SELECTORS)).all_text_contents()

        if len(info_items) >= 3:
            info_dict.update({
                "location": info_items[0],
                "experience": info_items[1],
                "education": info_items[2]
            })
        elif info_items:
            combined_text = " ".join(info_items)
            info_dict.update(self.parse_combined_info(combined_text))

        company_info = await self.safe_text(page, COMPANY_INFO_SELECTORS)
        if company_info:
            company_parsed = self.parse_company_info(company_info)
            info_dict.update(company_parsed)

        return info_dict

    def parse_combined_info(self, text):
        info = {}
        location_match = re.search(LOCATION_PATTERN, text)
        if location_match:
            info["location"] = location_match.group(1)

        for pattern in EXPERIENCE_PATTERNS:
            exp_match = re.search(pattern, text)
            if exp_match:
                info["experience"] = exp_match.group(1)
                break

        edu_match = re.search(EDUCATION_PATTERN, text)
        if edu_match:
            info["education"] = edu_match.group(1)
        return info

    def parse_company_info(self, text):
        info = {}
        size_match = re.search(COMPANY_SIZE_PATTERN, text)
        if size_match:
            info["company_size"] = size_match.group(1)
        return info

    async def parse_job_description(self, detail_raw):
        if not detail_raw:
            return "", ""
        cleaned_text = detail_raw.replace("\n", "\n").replace("\r", "").strip()

        for pattern in JOB_DESC_PATTERNS:
            match = re.search(pattern, cleaned_text, re.DOTALL | re.IGNORECASE)
            if match:
                groups = match.groups()
                if '职责' in groups[0] or '工作' in groups[0] or '岗位' in groups[2]:
                    return groups[1].strip(), groups[3].strip()
                else:
                    return groups[3].strip(), groups[1].strip()

        if '任职要求' in cleaned_text:
            parts = cleaned_text.split('任职要求')
            if len(parts) >= 2:
                duty = parts[0].replace('职责描述', '').replace('岗位职责', '').strip()
                require = parts[1].strip()
                return duty, require

        if '岗位职责' in cleaned_text or '工作职责' in cleaned_text:
            parts = re.split(r'[岗位职责工作职责职责描述]+', cleaned_text)
            if len(parts) >= 2:
                duty = parts[1].strip()
                require = parts[0].split('任职要求')[-1].strip() if '任职要求' in parts[0] else ""
                return duty, require

        return cleaned_text, ""

    async def get_job_links(self, page, url):
        for selector in JOB_LINK_SELECTORS:
            items = page.locator(selector)
            if await items.count() > 0:
                links = []
                for idx in range(await items.count()):
                    href = await items.nth(idx).get_attribute("href")
                    if href:
                        full_href = href if href.startswith("http") else "https://www.zhaopin.com" + href
                        links.append(full_href)
                return links
        return []

    async def scrape_job_detail(self, browser, url):
        page = await browser.new_page()
        try:
            await page.goto(url, wait_until="domcontentloaded", timeout=30000)
            await asyncio.sleep(random.uniform(SLEEP_MIN, SLEEP_MAX))

            job_data = {
                "岗位名称": await self.safe_text(page, JOB_TITLE_SELECTORS),
                "薪资": await self.safe_text(page, SALARY_SELECTORS),
                "公司名称": await self.safe_text(page, COMPANY_NAME_SELECTORS)
            }

            info_items = await self.get_basic_info(page)
            job_data.update({
                "工作地点": info_items.get("location", ""),
                "经验要求": info_items.get("experience", ""),
                "学历要求": info_items.get("education", ""),
                "公司规模": info_items.get("company_size", ""),
                "公司行业": info_items.get("company_industry", "")
            })

            detail_raw = await self.safe_text(page, DETAIL_CONTENT_SELECTORS)
            address = await self.safe_text(page, ADDRESS_SELECTORS)
            if address:
                job_data["工作地点"] = address

            duty, require = await self.parse_job_description(detail_raw)
            job_data.update({"工作职责": duty, "任职要求": require})
            job_data["详情链接"] = url
            job_data["抓取时间"] = asyncio.get_event_loop().time()

            return job_data

        except Exception as e:
            print(f"Error scraping {url}: {str(e)}")
            return {
                "岗位名称": "获取失败",
                "薪资": "获取失败",
                "工作地点": "获取失败",
                "经验要求": "获取失败",
                "学历要求": "获取失败",
                "工作职责": "获取失败",
                "任职要求": "获取失败",
                "详情链接": url,
                "抓取时间": asyncio.get_event_loop().time()
            }
        finally:
            await page.close()

    async def scrape_all_pages(self):
        all_new_jobs = []

        async with async_playwright() as p:
            browser_args = {}
            if PROXY:
                browser_args["proxy"] = PROXY

            browser = await p.chromium.launch(
                headless=False,
                args=[
                    '--no-sandbox',
                    '--disable-blink-features=AutomationControlled',
                    '--disable-dev-shm-usage'
                ],
                **browser_args
            )

            page = await browser.new_page()
            await page.set_extra_http_headers({
                "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
                "Accept-Encoding": "gzip, deflate, br",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8"
            })

            last_completed = max(self.crawled_pages) if self.crawled_pages else 0
            effective_start = max(START_PAGE, last_completed + 1)

            print(f" 配置起始页: {START_PAGE}, 已完成最大页: {last_completed}, 实际从第 {effective_start} 页开始")

            for i in range(effective_start, self.max_pages + 1):
                if i in self.crawled_pages:
                    continue

                url = re.sub(r"/p\d+", f"/p{i}", self.start_url)
                print(f"正在爬取第 {i} 页...")

                try:
                    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                    await page.wait_for_timeout(2000)

                    raw_links = await self.get_job_links(page, url)
                    if not raw_links:
                        print(f"第{i}页未找到职位链接")
                        self.crawled_pages.add(i)
                        self.save_resume_state()
                        continue

                    new_links = [link for link in raw_links if link not in self.seen_links]
                    print(f"第{i}页原始链接 {len(raw_links)} 个，去重后 {len(new_links)} 个新链接")

                    if not new_links:
                        print(f"第{i}页无新岗位，跳过")
                        self.crawled_pages.add(i)
                        self.save_resume_state()
                        continue

                    semaphore = asyncio.Semaphore(MAX_CONCURRENT)

                    async def sem_task(link):
                        async with semaphore:
                            return await self.scrape_job_detail(browser, link)

                    tasks = [sem_task(link) for link in new_links]
                    page_jobs = []

                    for f in tqdm_asyncio.as_completed(tasks, total=len(tasks), desc=f"第{i}页岗位", leave=False):
                        job = await f
                        if job and job.get("岗位名称") != "获取失败":
                            page_jobs.append(job)
                            self.seen_links.add(job["详情链接"])

                    all_new_jobs.extend(page_jobs)
                    print(f"第{i}页成功获取 {len(page_jobs)} 个新职位")

                    if page_jobs:
                        self.append_page_to_word(page_jobs, i)

                    self.crawled_pages.add(i)
                    self.save_resume_state()

                except Exception as e:
                    print(f"爬取第{i}页时出错: {str(e)}")
                    continue

            await browser.close()

        self.session_data.extend(all_new_jobs)
        return self.session_data

    def save_to_json(self, jobs, filename=JSON_FILENAME):
        seen = set()
        unique_jobs = []
        for job in jobs:
            link = job.get("详情链接")
            if link and link not in seen:
                unique_jobs.append(job)
                seen.add(link)
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(unique_jobs, f, ensure_ascii=False, indent=2)
        print(f"已保存完整去重数据到：{filename}")

    def print_summary(self, jobs):
        print(f"\n爬取完成！")
        print(f"总共获取到 {len(jobs)} 个有效职位（含历史）")
        if jobs:
            salaries = [job.get('薪资', '') for job in jobs if job.get('薪资')]
            locations = [job.get('工作地点', '') for job in jobs if job.get('工作地点')]
            print(f"薪资范围: {len([s for s in salaries if s])} 个职位有薪资信息")
            print(f"地点分布: {len([l for l in locations if l])} 个职位有地点信息")
            print(f"\n前3个职位预览:")
            for i, job in enumerate(jobs[:3], 1):
                print(f"{i}. {job.get('岗位名称', '未知')} - {job.get('薪资', '面议')} - {job.get('工作地点', '未知')}")


async def main():
    scraper = ZhaopinScraper(START_URL, MAX_PAGES)
    jobs = await scraper.scrape_all_pages()

    if jobs:
        scraper.save_to_json(jobs)
        scraper.print_summary(jobs)
    else:
        print("未能获取到任何职位信息")


if __name__ == "__main__":
    asyncio.run(main())